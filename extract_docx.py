"""Extract text and tables from all DEEVUH .docx specification files."""
import os
import sys

try:
    from docx import Document
except ImportError:
    os.system("pip install python-docx")
    from docx import Document

DOCX_DIR = os.path.join(os.path.dirname(__file__), "docx")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "docx_extracted")
os.makedirs(OUTPUT_DIR, exist_ok=True)

files = [
    "DEEVUH_ERD_SCHEMA_SPEC_V3.docx",
    "DEEVUH_SYSTEM_ARCHITECTURE_V3.docx",
    "DEEVUH_TRD_V3.docx",
    "DEEVUH_API_ROUTING_INTEGRATION_MAP_V3.docx",
    "DEEVUH_PRD_V3.docx",
    "DEEVUH_APP_FLOW_V3.docx",
    "DEEVUH_STATE_MACHINE_EDGE_CASE_SPEC_V3.docx",
    "DEEVUH_Deployment_Infrastructure_Multi_Region_Env_Mapping_V1.docx",
    "DEEVUH_V3_Project_Configuration_Parallel_Agent_Routing_Spec_Doc10.docx",
    "Deevuh_V3_Production_Implementation_Runbook.docx",
    "DEEVUH_Sandbox_Mocking_Synthetic_Data_Specification_V1.docx",
    "DEEVUH_Security_Compliance_Guardrail_Policy_V1.docx",
    "DEEVUH_Graceful_Degradation_High_Availability_Fallback_Playbook_V1.docx",
]

for fname in files:
    fpath = os.path.join(DOCX_DIR, fname)
    if not os.path.exists(fpath):
        print(f"SKIP: {fname} not found")
        continue
    
    doc = Document(fpath)
    out_name = fname.replace(".docx", ".txt")
    out_path = os.path.join(OUTPUT_DIR, out_name)
    
    with open(out_path, "w", encoding="utf-8") as f:
        # Extract paragraphs
        for para in doc.paragraphs:
            f.write(para.text + "\n")
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            f.write(f"\n--- TABLE {i+1} ---\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(cells) + "\n")
    
    print(f"OK: {fname} -> {out_name}")

print("\nDONE. All files extracted to docx_extracted/")
